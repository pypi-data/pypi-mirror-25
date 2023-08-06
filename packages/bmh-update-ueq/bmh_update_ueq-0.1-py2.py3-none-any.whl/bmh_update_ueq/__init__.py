# Author: Dan Levin
# Last changed: 22 Mar 17
# Dependancies: Python (2.7 is fine, 3.4 is better); python-docx; pandas
# Configuration instructions: set the config lines below to point to
# csvPath - file exported from wordpress
# docPath - folder containing .docx files, each of which should be named after the unit
# yearOfDocuments e.g. 2015-16
#
# If you have several years to import, do the earliest year, then change all 3 settings
# and then run again (using the output csv from the first batch as input for second)
#
# PDFs will need to be manually converted to .docx's
import sys
reload(sys)
sys.setdefaultencoding("utf-8") #needed on python 2.x
from docx import Document
from os import listdir
from os.path import isfile, join
import re
import codecs
import io
import pandas as pd
import time
import datetime
pd.set_option('display.max_colwidth', 100)

#Configure these lines. Be careful to give this script a single directory of .docx files from only ONE year (e.g 2015-16)

#docPath = '/Users/mqbssdl9/Documents/work/py/2015-16/' #MUST have trailing slash
#csvPath = '/Users/mqbssdl9/Documents/work/py/test.csv'
#yearOfDocuments = '2015-16'

docPath = '/Users/mqbssdl9/Documents/work/py/2016-17_mar17/' #MUST have trailing slash
csvPath = '/Users/mqbssdl9/Documents/work/py/test_2015-16_new.csv'
yearOfDocuments = '2016-17'
#End of Configuration

yearOfDocs = '<h1>' + yearOfDocuments + '</h1><br/>'



ts = time.time()
dateStamp = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')

fileList = [f for f in listdir(docPath) if isfile(join(docPath,f))]


#process csv to fix unicode in csv?
with codecs.open(csvPath,'r', 'utf-8') as f:
    text = f.read()
# process Unicode text
with codecs.open(csvPath,'w', 'utf-8') as f:
    f.write(text)


csv1 = pd.read_csv(csvPath,encoding="utf-8")
InCSV = csv1['post_tag'].tolist()

loc = [0,0,0,0] #location in file, for debug only

headers = ['Please highlight areas of good practic']
headers.append('Please highlight areas for improvement')
headers.append('What formative assessment was used')
headers.append('Please give an action plan of any changes')
headers.append('Unit Coordinator')
headers.append('Please list the lecturers who contribute')

data = {}
titles = {}

validFileCount = 0
CSVEntriesUpdated = 0
CSVEntriesAdded = 0

for f in fileList:
    print('NEWFILE: ' + str(f))
    if not f.endswith('.docx'):
        print('skipping invalid file')
    code = ''
    if f.endswith('.docx'):
        validFileCount += 1
        fName = f[0:-5].upper()
        if len(fName) == 9: #if there's only one code in the doc name, use that
            code = fName
            print('file contained 1 code: ' + code)

        if len(fName) > 9:
            print('long filename found; processing :' + fName)
            #find valid codes within the string
            pattern = re.compile(r"[A-Za-z]{4}\d{5}") #valid codes e.g. ABCD12345
            pattern2 = re.compile(r"[A-Za-z]{4}.*?[&| ]+?[0-9]{5}") #will match the closest 4 letters to the next 5 numbers
            codes=pattern.findall(fName)
            partCodes=pattern2.findall(fName)
            if len(partCodes) != 0:
                for c in partCodes:

                    d = c[0:4] + c[-5:] #each match should start with 4 letters and end with 5 numbers - stitch these to make a code
                    print ('code assembled: ' + d)
                    codes.append(d)



            for c in codes: #search spreadsheet for matching codes
                if code != '':
                    break
                for t in InCSV:
                    if code != '':
                        break
                    if c==t:
                        code = c
                        print('file contained multiple codes; code found with match: ' + code)
            if code == '':
                code = fName[0:9]
                print('file contained multiple codes, but none matched csv file. Code used: ' + code)



        loc = [0,0,0,0]
        doc = Document(docPath + f)

        tables = doc.tables
        entryText = yearOfDocs
        for table in tables:
            loc[0] += 1
            for row in table.rows:
                loc[1] += 1
                for cell in row.cells:
                    loc[2] += 1
                    for pgh in cell.paragraphs:
                        loc[3] += 1
                        #print(str(loc) + pgh.text)
                        if loc[0] == 2 and (loc[1] == 4 or loc[1] == 5 or (loc[1] == 6 and loc[2] == 10)):
                            if pgh.text != '':
                                foundText = pgh.text

                                for h in headers:

                                    if foundText.find(h) != -1:
                                        foundText = '<br/><b>' + foundText + '</b><br/>'
                                foundText += '<br/>'
                                #print(str(loc) + foundText)
                                entryText += foundText
                        elif loc[0] == 2 and loc[1] == 3 and loc[2] == 5 and loc[3] == 7:
                            titles[code] = pgh.text

        #print(entryText)

        data[code] = entryText #add key/val pair of code, content

NotInDocs = {}
tempList = sorted(list(InCSV))
NotInDocs = dict(enumerate(tempList))


for i, val in NotInDocs.items(): #NotInDocs begins with all CSV tags. each time one of these matches with a doc entry, remove it
    for k, v in data.items():
        if val == k:
            NotInDocs.pop(i, None)
#now NotInDocs contains just units that aren't being updated
if len(NotInDocs) > 0:
    print ('Units in the master spreadsheet but not in the datapack; for info only:')
    for i,val in NotInDocs.items():
        print (str(i) + ' ' + str(val))

NotInCSV = dict(data)

print()
print('Begin CSV Updates')
print()
for k,v in data.items(): #for each of the update pairs gathered
    for i, val in enumerate(InCSV):
        if str(val).upper() == str(k).upper(): #if the key from data matches
            #append the new info
            matchedEntry = csv1.loc[csv1['post_tag'] == k, 'post_content']
            print('matchFound: '+val+' ; B4Ud8: ' + matchedEntry)
            csv1.loc[csv1['post_tag'] == k, ['post_content']] = v + matchedEntry
            print('UpdatedCSV - code:' + k +': ' + csv1.loc[csv1['post_tag'] == k, 'post_content'])
            # make NotInCSV a list of incoming data which doesn't match rows in the CSV
            CSVEntriesUpdated += 1
            NotInCSV.pop(k, None)

    #print(str(k) + ': ' + str(v))
    #print(csv1.loc[csv1['post_tag'] == str(k), 'post_content'])
        #else:
            #print ('no match; val: ' + val + ' ; k: '+ k)
print('')
print('end of CSV updates')
print('')

if len(NotInCSV) != 0:
    print('')
    print ('Beginning CSV Inserts: ')
    print('')

    for k, v in NotInCSV.items():
        print('new item: ')
        print(k,titles[k],v[0:40])

        category = ''
        if re.search('[a-zA-Z]1',k):
            category = 'First Year'
        elif re.search('[a-zA-Z]2',k):
            category = 'Second Year'
        elif re.search('[a-zA-Z]3',k):
            category = 'Third Year'
        elif re.search('[a-zA-Z]4',k):
            category = '4th'
        elif re.search('[a-zA-Z]5',k):
            category = '5th'
        elif re.search('[a-zA-Z]6',k):
            category = '6th'

        codeAndName = k + ' ' + titles[k]
        csv1 = csv1.append({'post_title':codeAndName,'post_content': v, 'post_date':dateStamp,'post_name':k,'post_author':'1','post_status':'publish','comment_status':'0','ping_status':'closed','spacious_page_layout': 'default_layout','post_category':category,'post_tag':k},ignore_index=True)
        CSVEntriesAdded += 1
        #TODO: Look to optimize by appending multiple changes at once?
        print(k + ' :'+ csv1.loc[csv1['post_tag'] == str(k), 'post_content'])
    print('')
    print('end of CSV inserts')
    print('')
else:
    print('no new codes found')
    print('')

#save csv
csv1.to_csv(csvPath[:-4] + '_' + yearOfDocuments + '_new.csv', encoding='utf-8')

print('summary')
print('FilesProcessed: ' + str(validFileCount))
print('CSV Entries Updated: ' + str(CSVEntriesUpdated))
print('CSV Entries Added: ' + str(CSVEntriesAdded))






